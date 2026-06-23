#!/usr/bin/env python3
"""Build IJMI submission supplementary files from existing result artifacts.

This script is intentionally offline and deterministic. It does not rerun model
inference or training; it only recomputes lightweight analyses from saved
per-item prediction CSVs and writes submission-ready supplemental files.
"""
from __future__ import annotations

import argparse
import csv
import html
import json
import math
import os
import re
import shutil
import sys
import zipfile
from collections import defaultdict
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

import numpy as np
import pandas as pd
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from scripts.phase5_aggregate import _coverage_at_risk, _ece, _parse_json_col  # noqa: E402
from src.calibrate import fit_temperature, negative_log_likelihood  # noqa: E402
from src.models.registry import MODEL_SPECS, resolve_lora_target_modules  # noqa: E402
from src.models.score import softmax  # noqa: E402


DATASETS = ["vqa_rad", "slake_en", "pathvqa"]
MODELS = ["qwen25vl", "internvl", "llavaov", "smolvlm", "medgemma", "huatuo"]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def parse_test_name(path: Path) -> tuple[str, str, str] | None:
    match = re.match(
        r"phase4_(.+?)_(zero_shot|vqa_rad|slake_en|pathvqa)_on_"
        r"(vqa_rad|slake_en|pathvqa)_test\.csv$",
        path.name,
    )
    if match is None:
        return None
    return match.group(1), match.group(2), match.group(3)


def scores_and_labels(df: pd.DataFrame) -> tuple[list[list[float]], list[int]]:
    scores = _parse_json_col(df["normalized_scores"])
    labels = df["gold_idx"].astype(int).tolist()
    keep = [(s, y) for s, y in zip(scores, labels, strict=True) if s and 0 <= y < len(s)]
    return [s for s, _ in keep], [y for _, y in keep]


def metrics_from_scores(
    scores: list[list[float]],
    labels: list[int],
    temperature: float,
) -> tuple[float, float, float]:
    confs: list[float] = []
    correct: list[float] = []
    for item_scores, gold in zip(scores, labels, strict=True):
        probs = softmax(item_scores, temperature=temperature)
        pred = int(np.argmax(probs))
        confs.append(float(max(probs)))
        correct.append(float(pred == gold))
    accuracy = float(np.mean(correct))
    ece = _ece(np.array(confs), np.array(correct))
    nll = negative_log_likelihood(scores, labels, temperature)
    return accuracy, ece, nll


def load_prediction(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def write_csv(df: pd.DataFrame, path: Path) -> None:
    df.to_csv(path, index=False, float_format="%.6f")


def make_source_vs_target_temperature(results_dir: Path, out_path: Path) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for model in MODELS:
        for train_dataset in DATASETS:
            source_calib = results_dir / f"phase5_calib_{model}_{train_dataset}_on_{train_dataset}_calib.csv"
            if not source_calib.exists():
                continue
            source_scores, source_labels = scores_and_labels(load_prediction(source_calib))
            if not source_labels:
                continue
            source_fit = fit_temperature(source_scores, source_labels)
            for eval_dataset in DATASETS:
                if eval_dataset == train_dataset:
                    continue
                target_calib = results_dir / f"phase5_calib_{model}_{train_dataset}_on_{eval_dataset}_calib.csv"
                test_path = results_dir / f"phase4_{model}_{train_dataset}_on_{eval_dataset}_test.csv"
                if not target_calib.exists() or not test_path.exists():
                    continue
                target_scores, target_labels = scores_and_labels(load_prediction(target_calib))
                test_scores, test_labels = scores_and_labels(load_prediction(test_path))
                if not target_labels or not test_labels:
                    continue
                target_fit = fit_temperature(target_scores, target_labels)
                acc_unscaled, ece_unscaled, nll_unscaled = metrics_from_scores(
                    test_scores, test_labels, 1.0
                )
                _, ece_source, nll_source = metrics_from_scores(
                    test_scores, test_labels, source_fit.temperature
                )
                _, ece_target, nll_target = metrics_from_scores(
                    test_scores, test_labels, target_fit.temperature
                )
                rows.append(
                    {
                        "model": model,
                        "train_dataset": train_dataset,
                        "eval_dataset": eval_dataset,
                        "n_source_calib": len(source_labels),
                        "n_target_calib": len(target_labels),
                        "n_test": len(test_labels),
                        "source_temperature": source_fit.temperature,
                        "target_temperature": target_fit.temperature,
                        "accuracy": acc_unscaled,
                        "ece_unscaled": ece_unscaled,
                        "ece_source_temperature": ece_source,
                        "ece_target_temperature": ece_target,
                        "target_minus_source_ece": ece_target - ece_source,
                        "nll_unscaled": nll_unscaled,
                        "nll_source_temperature": nll_source,
                        "nll_target_temperature": nll_target,
                        "target_minus_source_nll": nll_target - nll_source,
                    }
                )
    out = pd.DataFrame(rows)
    write_csv(out, out_path)
    return out


def make_selective_review(results_dir: Path, out_path: Path) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for path in sorted(results_dir.glob("phase4_*_on_*_test.csv")):
        if any(suffix in path.name for suffix in ["_s1", "_s2", "_s3"]):
            continue
        parsed = parse_test_name(path)
        if parsed is None:
            continue
        model, condition, eval_dataset = parsed
        df = load_prediction(path)
        confs = df["confidence"].astype(float).values
        correct = df["correct"].astype(int).values.astype(float)
        n = len(df)
        error = correct == 0
        high = confs >= 0.90
        row: dict[str, object] = {
            "model": model,
            "condition": condition,
            "eval_dataset": eval_dataset,
            "evaluation_type": (
                "zero_shot"
                if condition == "zero_shot"
                else "fine_tuned_id"
                if condition == eval_dataset
                else "fine_tuned_cross_dataset"
            ),
            "n": n,
            "accuracy_all": float(correct.mean()),
            "error_n": int(error.sum()),
            "high_conf_error_n": int((error & high).sum()),
            "high_conf_error_pct_of_errors": (
                float((error & high).sum() / error.sum()) if error.sum() else math.nan
            ),
            "high_conf_error_pct_of_all_items": float((error & high).sum() / n),
            "coverage_at_risk_le_5pct": _coverage_at_risk(confs, correct, max_risk=0.05),
        }
        order = np.argsort(-confs)
        for coverage in [0.9, 0.8, 0.7, 0.5]:
            k = max(1, int(math.ceil(coverage * n)))
            kept = correct[order[:k]]
            row[f"selective_accuracy_at_{int(coverage * 100)}pct_coverage"] = float(kept.mean())
            row[f"risk_at_{int(coverage * 100)}pct_coverage"] = float(1.0 - kept.mean())
            row[f"review_fraction_at_{int(coverage * 100)}pct_coverage"] = float(1.0 - coverage)
        rows.append(row)
    out = pd.DataFrame(rows)
    write_csv(out, out_path)
    return out


def resolved_model_table(results_dir: Path) -> pd.DataFrame:
    resolved: dict[str, dict[str, str]] = {}
    for path in sorted(results_dir.glob("phase4_*_on_*_test.csv")):
        parsed = parse_test_name(path)
        if parsed is None:
            continue
        model = parsed[0]
        if model in resolved:
            continue
        df = pd.read_csv(path, nrows=1)
        resolved[model] = {
            "resolved_hf_id": str(df.get("resolved_hf_id", pd.Series([""])).iloc[0]),
            "model_family_runtime": str(df.get("model_family", pd.Series([""])).iloc[0]),
            "model_type_runtime": str(df.get("model_type", pd.Series([""])).iloc[0]),
        }
    rows = []
    for model in MODELS:
        spec = MODEL_SPECS[model]
        rows.append(
            {
                "model": model,
                "configured_hf_id": spec.hf_id,
                "resolved_hf_id": resolved.get(model, {}).get("resolved_hf_id", ""),
                "family_configured": spec.family,
                "family_runtime": resolved.get(model, {}).get("model_family_runtime", ""),
                "model_type_configured": spec.model_type,
                "model_type_runtime": resolved.get(model, {}).get("model_type_runtime", ""),
                "size": spec.size,
                "gated": spec.gated,
                "lora_target_modules": ", ".join(resolve_lora_target_modules(spec)),
            }
        )
    return pd.DataFrame(rows)


def make_weighting_dedup_sensitivity(
    master_metrics: pd.DataFrame,
    model_table: pd.DataFrame,
    out_path: Path,
) -> pd.DataFrame:
    metric_cols = ["accuracy", "ece", "nll", "aurc"]
    model_to_resolved = dict(zip(model_table["model"], model_table["resolved_hf_id"], strict=False))
    df = master_metrics.copy()
    df["resolved_hf_id"] = df["model"].map(model_to_resolved).fillna(df["model"])
    rows: list[dict[str, object]] = []
    for condition_group, group_df in [
        ("zero_shot", df[df["condition"] == "zero_shot"]),
        ("fine_tuned_id", df[(df["condition"] != "zero_shot") & (df["condition"] == df["eval_dataset"])]),
        (
            "fine_tuned_cross_dataset",
            df[(df["condition"] != "zero_shot") & (df["condition"] != df["eval_dataset"])],
        ),
        ("all_cells", df),
    ]:
        if group_df.empty:
            continue
        for metric in metric_cols:
            by_checkpoint = (
                group_df.groupby(["resolved_hf_id", "condition", "eval_dataset"], dropna=False)[metric]
                .mean()
                .reset_index()
            )
            rows.append(
                {
                    "condition_group": condition_group,
                    "metric": metric,
                    "n_cells": len(group_df),
                    "macro_mean": float(group_df[metric].mean()),
                    "sample_weighted_mean": float(
                        np.average(group_df[metric].astype(float), weights=group_df["n"].astype(float))
                    ),
                    "resolved_checkpoint_macro_mean": float(by_checkpoint[metric].mean()),
                    "n_resolved_checkpoints": int(group_df["resolved_hf_id"].nunique()),
                }
            )
    out = pd.DataFrame(rows)
    write_csv(out, out_path)
    return out


def summarize_rows(df: pd.DataFrame, group_col: str, value_cols: list[str]) -> pd.DataFrame:
    rows = []
    for key, g in df.groupby(group_col, dropna=False):
        row: dict[str, object] = {group_col: key, "n_rows": len(g)}
        for col in value_cols:
            if col in g:
                row[f"mean_{col}"] = float(g[col].mean())
        rows.append(row)
    return pd.DataFrame(rows)


def xml_cell(value: object, row_idx: int, col_idx: int) -> str:
    col = ""
    idx = col_idx
    while True:
        idx, rem = divmod(idx, 26)
        col = chr(ord("A") + rem) + col
        if idx == 0:
            break
        idx -= 1
    ref = f"{col}{row_idx}"
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return f'<c r="{ref}"/>'
    if isinstance(value, (int, np.integer)) and not isinstance(value, bool):
        return f'<c r="{ref}"><v>{int(value)}</v></c>'
    if isinstance(value, (float, np.floating)):
        return f'<c r="{ref}"><v>{float(value):.10g}</v></c>'
    text = escape(str(value))
    return f'<c r="{ref}" t="inlineStr"><is><t>{text}</t></is></c>'


def worksheet_xml(df: pd.DataFrame) -> str:
    rows = []
    headers = list(df.columns)
    rows.append(
        '<row r="1">' + "".join(xml_cell(header, 1, i) for i, header in enumerate(headers)) + "</row>"
    )
    for r, (_, row) in enumerate(df.iterrows(), start=2):
        rows.append(
            f'<row r="{r}">'
            + "".join(xml_cell(row[col], r, i) for i, col in enumerate(headers))
            + "</row>"
        )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheetData>'
        + "".join(rows)
        + "</sheetData></worksheet>"
    )


def make_xlsx(sheets: dict[str, pd.DataFrame], out_path: Path) -> None:
    sheet_items = list(sheets.items())
    content_types = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">',
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="xml" ContentType="application/xml"/>',
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>',
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
    ]
    for idx in range(1, len(sheet_items) + 1):
        content_types.append(
            f'<Override PartName="/xl/worksheets/sheet{idx}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        )
    content_types.append("</Types>")

    workbook_sheets = []
    workbook_rels = []
    for idx, (name, _) in enumerate(sheet_items, start=1):
        safe_name = html.escape(name[:31], quote=True)
        workbook_sheets.append(f'<sheet name="{safe_name}" sheetId="{idx}" r:id="rId{idx}"/>')
        workbook_rels.append(
            f'<Relationship Id="rId{idx}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
            f'Target="worksheets/sheet{idx}.xml"/>'
        )
    workbook_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "<sheets>"
        + "".join(workbook_sheets)
        + "</sheets></workbook>"
    )
    rels_root = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        '<Relationship Id="rId2" '
        'Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" '
        'Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" '
        'Target="docProps/app.xml"/>'
        "</Relationships>"
    )
    workbook_rels_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(workbook_rels)
        + "</Relationships>"
    )
    app_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties">'
        "<Application>Python</Application></Properties>"
    )
    core_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:dcterms="http://purl.org/dc/terms/" '
        'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
        'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        "<dc:title>Supplementary tables</dc:title></cp:coreProperties>"
    )
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", "".join(content_types))
        zf.writestr("_rels/.rels", rels_root)
        zf.writestr("docProps/app.xml", app_xml)
        zf.writestr("docProps/core.xml", core_xml)
        zf.writestr("xl/workbook.xml", workbook_xml)
        zf.writestr("xl/_rels/workbook.xml.rels", workbook_rels_xml)
        for idx, (_, df) in enumerate(sheet_items, start=1):
            zf.writestr(f"xl/worksheets/sheet{idx}.xml", worksheet_xml(df))


def add_docx_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Status"
    hdr[2].text = "Where addressed"
    for item, status, where in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = status
        cells[2].text = where


def make_ml_checklist(out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Machine-learning reporting checklist for IJMI submission", level=0)
    doc.add_paragraph(
        "This checklist summarizes information provided in the manuscript and supplementary "
        "appendix for an empirical evaluation of confidence reliability in fine-tuned medical "
        "vision-language models."
    )
    rows = [
        ("Clinical or workflow rationale", "Reported", "Introduction; Discussion; Supplement A.1"),
        ("Datasets, source, inclusion/exclusion criteria", "Reported", "Methods; Table 1; Supplement A.2; Table S1"),
        ("Train/calibration/test separation", "Reported", "Methods; Supplement A.2; Table S1"),
        ("External or cross-dataset validation", "Reported", "Methods and Results RQ4; Table S3; Table S9"),
        ("Model identifiers and versions", "Reported", "Table 2; Table S2"),
        ("Preprocessing and option construction", "Reported", "Methods; Supplement A.3"),
        ("Training recipe and hyperparameters", "Reported", "Table 3; Supplement A.4"),
        ("Hardware and software environment", "Reported", "Supplement A.10"),
        ("Outcome metrics and uncertainty intervals", "Reported", "Methods; Table S3"),
        ("Calibration and selective prediction", "Reported", "Methods; Results RQ2/RQ3; Tables S4/S8/S9"),
        ("Handling of distribution shift", "Reported", "Results RQ4/RQ5; Tables S5/S7/S9"),
        ("Biases and limitations", "Reported", "Limitations; Supplement A.11"),
        ("Data and code availability", "Prepared", "Supplementary data/code availability statement"),
        ("Ethics and privacy", "Reported", "Ethics, reporting, and reproducibility"),
    ]
    add_docx_table(doc, rows)
    doc.add_paragraph(
        "No new model training or inference was performed for the supplementary analyses; they "
        "are recomputations from saved per-item prediction files."
    )
    doc.save(out_path)


def make_data_code_statement(out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Data and code availability statement", level=0)
    doc.add_paragraph(
        "All primary datasets used in this study are public benchmark datasets. The planned "
        "reproducibility release should include split identifiers, prompts and answer-option "
        "construction, raw per-item prediction CSV files, metric aggregation scripts, temperature "
        "scaling scripts, corruption definitions, software versions, resolved model identifiers, "
        "and dataset license information."
    )
    doc.add_paragraph(
        "Large model weights are not redistributed. Public model checkpoints can be obtained from "
        "their original providers subject to each model's license and access restrictions."
    )
    doc.add_paragraph(
        "Before submission, replace this paragraph with the final repository or archive DOI "
        "(for example Zenodo, Mendeley Data, OSF, or an institutional repository)."
    )
    doc.save(out_path)


def make_highlights(out_path: Path) -> None:
    highlights = [
        "Fine-tuning improved accuracy and average calibration but heterogeneously.",
        "Medical model labelling was not a sufficient calibration proxy.",
        "Option-likelihood confidence outperformed verbalized self-confidence.",
        "Target-domain temperature scaling reduced miscalibration.",
        "Cross-dataset and modality shifts were larger stressors than corruptions.",
    ]
    doc = Document()
    doc.add_heading("Highlights", level=0)
    for item in highlights:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(out_path)


def make_summary_table(out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Summary Table", level=0)
    doc.add_heading("What was already known", level=1)
    known = [
        (
            "Vision-language models are increasingly adapted to medical tasks with "
            "parameter-efficient fine-tuning, and fine-tuning is widely assumed to "
            "increase overconfidence."
        ),
        (
            "Calibration and selective prediction are relevant to the reliability of "
            "adapted medical VLMs, yet are rarely reported."
        ),
    ]
    for item in known:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("What this study adds", level=1)
    added = [
        (
            "Across six models and three datasets, fine-tuning improved in-distribution "
            "accuracy and average calibration on average, but heterogeneously; medical "
            "pre-training did not guarantee calibrated confidence."
        ),
        (
            "Cross-dataset and cross-modality distribution shift, not image corruption, "
            "was the main cause of unreliable confidence; target-domain temperature "
            "scaling reduced miscalibration but required labelled calibration data."
        ),
    ]
    for item in added:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(out_path)


def latex_escape(text: object) -> str:
    s = str(text)
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(ch, ch) for ch in s)


def df_to_latex_tabular(df: pd.DataFrame, columns: list[str], max_rows: int = 12) -> str:
    shown = df[columns].head(max_rows)
    col_spec = "l" * len(columns)
    lines = [r"\noindent\resizebox{\linewidth}{!}{%", rf"\begin{{tabular}}{{{col_spec}}}", r"\toprule"]
    lines.append(" & ".join(latex_escape(c) for c in columns) + r" \\")
    lines.append(r"\midrule")
    for _, row in shown.iterrows():
        vals = []
        for c in columns:
            val = row[c]
            if isinstance(val, float):
                vals.append(f"{val:.3f}")
            else:
                vals.append(latex_escape(val))
        lines.append(" & ".join(vals) + r" \\")
    lines.extend([r"\bottomrule", r"\end{tabular}%", r"}"])
    return "\n".join(lines)


def make_appendix_tex(
    out_path: Path,
    dataset_table: pd.DataFrame,
    model_table: pd.DataFrame,
    master_metrics: pd.DataFrame,
    source_target: pd.DataFrame,
    selective_review: pd.DataFrame,
    weighted: pd.DataFrame,
) -> None:
    id_cells = master_metrics[
        (master_metrics["condition"] != "zero_shot")
        & (master_metrics["condition"] == master_metrics["eval_dataset"])
    ]
    clean_summary = summarize_rows(master_metrics, "condition", ["accuracy", "ece", "nll", "aurc"])
    selective_summary = summarize_rows(
        selective_review,
        "evaluation_type",
        [
            "accuracy_all",
            "coverage_at_risk_le_5pct",
            "high_conf_error_pct_of_errors",
            "selective_accuracy_at_70pct_coverage",
        ],
    )
    st_summary = pd.DataFrame(
        [
            {
                "n_cross_cells": len(source_target),
                "mean_ece_unscaled": source_target["ece_unscaled"].mean(),
                "mean_ece_source_temperature": source_target["ece_source_temperature"].mean(),
                "mean_ece_target_temperature": source_target["ece_target_temperature"].mean(),
                "target_better_ece_cells": int(
                    (source_target["ece_target_temperature"] < source_target["ece_source_temperature"]).sum()
                ),
            }
        ]
    )
    text = rf"""\documentclass[11pt]{{article}}
\usepackage[margin=1in]{{geometry}}
\usepackage{{booktabs}}
\usepackage{{graphicx}}
\usepackage{{longtable}}
\usepackage{{array}}
\usepackage{{hyperref}}
\title{{Supplementary Appendix: Calibration and Selective Prediction of Fine-Tuned Medical VLMs}}
\date{{}}
\begin{{document}}
\maketitle

\section*{{A.1 Study scope and reproducibility rationale}}
This appendix provides the transparency material requested for an AI/ML study in
International Journal of Medical Informatics. The study is an empirical audit of
confidence reliability after parameter-efficient fine-tuning. It does not
introduce a new model architecture. All supplementary analyses are recomputed
from saved per-item prediction CSV files rather than from additional model
training or inference.

\section*{{A.2 Dataset construction and splits}}
Closed-ended items were retained so that correctness and option likelihoods were
well defined. Open-ended free-text answers were excluded from the core analysis.
Splits used fixed seed 42 and separated training, calibration, and test data.

{df_to_latex_tabular(dataset_table, ["dataset", "raw_count", "kept_count", "split", "split_count"], 15)}

\section*{{A.3 Model registry and resolved checkpoints}}
Every prediction row stores the model registry name, resolved checkpoint,
model-family label, and medical/general label.

{df_to_latex_tabular(model_table, ["model", "resolved_hf_id", "model_type_runtime", "size"], 10)}

\section*{{A.4 Fine-tuning configuration}}
All adapters used QLoRA with a frozen vision encoder, rank 16, alpha 32,
dropout 0.05, batch size 4 with gradient accumulation 4, maximum sequence length
1024, bfloat16 where supported, and fixed seed 42. Architecture-specific
language-side projection names were resolved through the model registry.

\section*{{A.5 Confidence scoring}}
For each answer option, the model was teacher-forced on the answer tokens under
the image-question prompt. The primary option score was the length-normalized
answer log likelihood. A softmax over option scores yielded the predicted answer
and confidence. Raw summed log likelihood was retained as an ablation input but
was not the primary confidence score.

\section*{{A.6 Full clean-grid metrics}}
Table S3 contains all 72 clean evaluation cells. Macro summaries are shown below;
tables and CSV files contain the complete values with confidence intervals.

{df_to_latex_tabular(clean_summary, ["condition", "n_rows", "mean_accuracy", "mean_ece", "mean_nll", "mean_aurc"], 12)}

\section*{{A.7 Post-hoc calibration}}
Temperature scaling was fit on a labelled calibration split and applied to the
corresponding test split. The main analysis uses the target evaluation dataset's
calibration split. As a sensitivity check, Table S9 compares cross-dataset
evaluation using temperatures fit on the source training dataset versus the
target evaluation dataset.

{df_to_latex_tabular(st_summary, ["n_cross_cells", "mean_ece_unscaled", "mean_ece_source_temperature", "mean_ece_target_temperature", "target_better_ece_cells"], 5)}

\section*{{A.8 Selective review operating points}}
Table S8 reports selective accuracy and risk at 90, 80, 70, and 50\% coverage,
coverage at risk $\leq 5\%$, and high-confidence error counts. Summary:

{df_to_latex_tabular(selective_summary, ["evaluation_type", "n_rows", "mean_accuracy_all", "mean_coverage_at_risk_le_5pct", "mean_high_conf_error_pct_of_errors", "mean_selective_accuracy_at_70pct_coverage"], 10)}

\section*{{A.9 Distribution shift and corruptions}}
Table S5 gives the complete corruption grid. Table S7 gives OmniMedVQA
modality-level results. These analyses separate same-dataset image degradation
from cross-dataset and cross-modality transfer.

\section*{{A.10 Weighting and resolved-checkpoint sensitivity}}
Main text summaries are unweighted macro-averages over model-dataset cells.
Table S10 reports sample-weighted means and grouping by resolved checkpoint.

{df_to_latex_tabular(weighted, ["condition_group", "metric", "macro_mean", "sample_weighted_mean", "resolved_checkpoint_macro_mean"], 12)}

\section*{{A.11 Limitations relevant to clinical translation}}
The study uses retrospective public benchmarks rather than a prospective
workflow deployment. Closed-ended VQA enables rigorous option-likelihood
confidence scoring, but does not cover open-ended report generation. Cross-dataset
transfer is a stress test of source shift and answer-prior shift, not a substitute
for hospital-level external validation. Temperature scaling requires labelled
calibration data from the target domain.

\section*{{Supplementary figures}}
\begin{{figure}}[h]
\centering
\includegraphics[width=0.95\linewidth]{{figures/lora_supp.pdf}}
\caption{{Supplementary Figure S1. Low-rank adaptation and QLoRA configuration.}}
\end{{figure}}

\end{{document}}
"""
    out_path.write_text(text, encoding="utf-8")


def copy_supplementary_figures(figures_dir: Path, out_dir: Path, paper_images: Path) -> None:
    ensure_dir(out_dir)
    for name in ["lora_supp.pdf"]:
        shutil.copy2(paper_images / name, out_dir / name)
    for path in sorted(figures_dir.glob("reliability_*.png")):
        shutil.copy2(path, out_dir / path.name)
    for name in [
        "fig1_rq1_ece_before_after.png",
        "fig2_risk_coverage_vqa_rad.png",
        "fig2_risk_coverage_slake_en.png",
        "fig2_risk_coverage_pathvqa.png",
        "fig3_corruption_degradation.png",
        "fig4_modality_calibration.png",
        "fig4b_aurc_modality_calibration.png",
    ]:
        src = figures_dir / name
        if src.exists():
            shutil.copy2(src, out_dir / name)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--repo-root", default=".")
    parser.add_argument("--paper-dir", default="../paper_overleaf")
    args = parser.parse_args()

    repo = Path(args.repo_root).resolve()
    paper = Path(args.paper_dir).resolve()
    results = repo / "results"
    out = paper / "supplementary"
    tables_dir = out / "tables"
    figures_dir = out / "figures"
    ensure_dir(out)
    ensure_dir(tables_dir)

    source_target = make_source_vs_target_temperature(
        results, results / "source_vs_target_temperature.csv"
    )
    selective_review = make_selective_review(
        results, results / "selective_review_operating_points.csv"
    )

    dataset_table = pd.read_csv(results / "dataset_counts.csv")
    model_table = resolved_model_table(results)
    master_metrics = pd.read_csv(results / "master_metrics.csv")
    temperature = pd.read_csv(results / "temperature_scaling_calib.csv")
    corruptions = pd.read_csv(results / "rq4_corruption_metrics.csv")
    verbalized = pd.read_csv(results / "rq2_verbalized_metrics.csv")
    modality = pd.read_csv(results / "rq5_modality_metrics.csv")
    weighted = make_weighting_dedup_sensitivity(
        master_metrics, model_table, results / "weighting_dedup_sensitivity.csv"
    )

    table_outputs: dict[str, pd.DataFrame] = {
        "S1_dataset_splits": dataset_table,
        "S2_model_registry": model_table,
        "S3_clean_metrics": master_metrics,
        "S4_target_temperature": temperature,
        "S5_corruptions": corruptions,
        "S6_verbalized": verbalized,
        "S7_modality": modality,
        "S8_selective_review": selective_review,
        "S9_source_vs_target_T": source_target,
        "S10_weighting_dedup": weighted,
    }
    for name, df in table_outputs.items():
        write_csv(df, tables_dir / f"{name}.csv")
    make_xlsx(table_outputs, out / "Supplementary_Tables.xlsx")

    make_ml_checklist(out / "IJMI_ML_Checklist.docx")
    make_data_code_statement(out / "Data_and_Code_Availability.docx")
    make_highlights(out / "Highlights.docx")
    make_summary_table(out / "Summary_Table.docx")
    copy_supplementary_figures(repo / "figures", figures_dir, paper / "images")
    make_appendix_tex(
        out / "Supplementary_Appendix.tex",
        dataset_table,
        model_table,
        master_metrics,
        source_target,
        selective_review,
        weighted,
    )
    readme = """# IJMI supplementary submission package

Generated by `scripts/39_make_submission_supplements.py`.

Files:

- `Supplementary_Appendix.tex` / `Supplementary_Appendix.pdf`: narrative appendix with methods, transparency notes, and summary tables.
- `Supplementary_Tables.xlsx`: multi-sheet workbook for Tables S1-S10.
- `tables/*.csv`: CSV versions of each supplementary table.
- `IJMI_ML_Checklist.docx`: completed machine-learning reporting checklist.
- `Data_and_Code_Availability.docx`: draft data/code availability statement; replace repository/DOI placeholder before final submission.
- `Highlights.docx`: separate editable highlights file.
- `Summary_Table.docx`: IJMI-style summary table with 2 already-known and 2 added-value bullets.
- `figures/`: copied supplementary figure assets.

No new model inference or training is run by this script.
"""
    (out / "README_submission_files.md").write_text(readme, encoding="utf-8")

    print(f"Wrote supplementary package to {out}")
    print(f"source-vs-target calibration rows: {len(source_target)}")
    print(f"selective review rows: {len(selective_review)}")


if __name__ == "__main__":
    main()
